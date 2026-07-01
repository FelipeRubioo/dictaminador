from docx import Document
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from datetime import datetime
from docx.shared import Cm, Pt, RGBColor
import config
from docx2pdf import convert


def get_spanish_month(month_number):
    months = [
        "enero", "febrero", "marzo", "abril", "mayo", "junio",
        "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
    ]
    return months[month_number - 1]

def agregarEncabezado(doc: Document,numeroDictamen,folio,anio):
    now = datetime.now()

    day = now.day
    month = get_spanish_month(now.month)
    year = now.year

    textoFecha = f"Hermosillo, Sonora, a {day} de {month} de {year}"

    logo_path = Path(__file__).parent / "images" / "logoSTJ.png"
    stacked_text = f"\tDictamen:{numeroDictamen}/{anio}\n\tReferencia a la solicitud de Soporte Técnico {folio}/{anio}\n\t{textoFecha}"
    for section in doc.sections:
    
      header = section.header
      paragraph1 = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

      run = paragraph1.add_run()
      run.add_picture(str(logo_path), width=Cm(2.5))
    
      paragraph2 = header.add_paragraph()
      paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

      run = paragraph2.add_run(stacked_text)
      run.bold = True
      run.font.size = Pt(10)

def agregarTitular(doc: Document,nombreTitular, puestoTitular,unidad):
    stacked_text = f"{nombreTitular}\n{puestoTitular}\n{unidad}\nPresente.-"
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(stacked_text)
    run.bold = True
    run.font.size = Pt(12)

def agregarIntroduccion(doc: Document,folio,anio):
    stacked_text = f"En atención a la solicitud de Soporte Técnico {folio}/{anio}, se emite el presente dictamen."
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(stacked_text)
    run.font.size = Pt(11)

def agregarTabla(doc: Document,solicitante,modelo,inventario,serie,fechaCompra):
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    # Fill in the table with data
    data = [
        ("Usuario:", solicitante),
        ("Modelo:", modelo),
        ("Número de Inventario:", inventario),
        ("Número de Serie:", serie),
        ("Fecha de Compra:", fechaCompra)
    ]

    for row, (label, value) in zip(table.rows, data):
        row.cells[0].text = label
        row.cells[1].text = value

        run = row.cells[1].paragraphs[0].runs[0]
        run.font.color.rgb = RGBColor(0, 0, 255)  # Blue

def agregarDiagnostico(doc: Document,diagnostico,imgDiagnosticoPath=""):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"\n{diagnostico}")
    run.font.size = Pt(11)
    if(len(imgDiagnosticoPath)>0):
        paragraph = doc.add_paragraph()
        run2 = paragraph.add_run()
        run2.add_picture(str(imgDiagnosticoPath),width=Cm(5))

def agregarConclusion(doc: Document,tipoDictamen):
    if tipoDictamen == "baja":
        conclusion_text = f"El equipo listado es candidato para REEMPLAZO, dando seguimiento al oficio de 'Requisitos mínimos para equipos de cómputo' con referencia DGSC-285/2023 del 17 de octubre del 2023, donde cito: 'De igual manera, se establece que, para que los equipos actuales sean considerados candidatos a actualización de componentes como la RAM y el Disco Duro, deben cumplir con las siguientes características: una placa madre que soporte al menos 16GB de RAM, un procesador Core i3 de octava generación en adelante, y un tiempo de adquisición no mayor a 5 años'."
    if tipoDictamen == "actualizacion":
        conclusion_text = f"El equipo listado es candidato para ACTUALIZACION, dando seguimiento al oficio de 'Requisitos mínimos para equipos de cómputo' con referencia DGSC-285/2023 del 17 de octubre del 2023, donde cito: 'De igual manera, se establece que, para que los equipos actuales sean considerados candidatos a actualización de componentes como la RAM y el Disco Duro, deben cumplir con las siguientes características: una placa madre que soporte al menos 16GB de RAM, un procesador Core i3 de octava generación en adelante, y un tiempo de adquisición no mayor a 5 años'."

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(conclusion_text)
    run.font.size = Pt(11)

def agregarRecomendacion(doc: Document,tipoDictamen,tipoBaja = "",componente="",linkCompra = ""):
    equipoNuevo = ""
    if tipoDictamen == "baja":
        if tipoBaja == "computadora":
            equipoNuevo = "Lenovo ThinkCentre neo 50a Gen 5 AIO "
        if tipoBaja == "laptop":
            equipoNuevo = "Dell Precision 3581 Laptop "
        if tipoBaja == "impresora":
            equipoNuevo = "Canon imagerunner 1643ii"
        if tipoBaja == "regulador":
            equipoNuevo = "No Break CDP R-UPR1008, 500W, 1000VA, 8 Contactos, Entrada 80-145V, Salida 120V"
        if tipoBaja == "monitor":
            equipoNuevo = "Samsung Monitor 24 FHD LS24F330EALXZX"
        recomendacion_text = f"Se recomienda sustituir el equipo listado por un {equipoNuevo} para garantizar un rendimiento óptimo y cumplir con los estándares de la institución."
    if tipoDictamen == "actualizacion":
        recomendacion_text = f"Se recomienda actualizar el equipo listado y comprar el componente: {componente}. Link de compra: {linkCompra}."

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(recomendacion_text)
    run.font.size = Pt(11)

def agregarRemitente(doc: Document):
    stacked_text = f"\t\t\t{config.ABREVIATURA} {config.NOMBRESTJ}\n\t\tSupremo Tribunal de Justicia del Estado de Sonora\n\tDirección General de Tecnologías de la Información y la Comunicación"
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(stacked_text)
    run.bold = True
    run.font.size = Pt(11)

def agregarOficio(doc: Document):
    text= "ANEXO OFICIO DGSC-285/2023"
    oficio_path = Path(__file__).parent / "images" / "oficio.png"

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.add_picture(str(oficio_path))


def generarDictamen(folio,anio,unidad,solicitante,inventario,serie,fechaCompra,nombreTitular, puestoTitular,numeroDictamen,modelo,tipoDictamen,diagnostico,imgDiagnosticoPath="",tipoBaja="",componente="",linkCompra=""):
    #creates file
    doc = Document()
    #set global style
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    #make all changes to file
    agregarEncabezado(doc,numeroDictamen,folio,anio)
    agregarTitular(doc,nombreTitular,puestoTitular,unidad)
    agregarIntroduccion(doc,folio,anio)
    agregarTabla(doc,solicitante,modelo,inventario,serie,fechaCompra)
    agregarDiagnostico(doc,diagnostico,imgDiagnosticoPath)
    agregarConclusion(doc,tipoDictamen)
    agregarRecomendacion(doc,tipoDictamen,tipoBaja,componente,linkCompra)
    agregarRemitente(doc)
    agregarOficio(doc)
    
    #crea directorio con el numero del ticket
    target_dir = Path(fr"{config.DIRECTORIO}\{folio}-{anio}")
    target_dir.mkdir(parents=True, exist_ok=True)
    #crea directorio con folio de ticket
    dictamen_path = target_dir / f"{folio}-{anio}.docx"
    doc.save(dictamen_path)
    convert(dictamen_path, target_dir / f"{folio}-{anio}.pdf")
    return target_dir / f"{folio}-{anio}.pdf"